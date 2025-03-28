import os
from pathlib import Path

# Test data directory
TEST_DATA_DIR = Path(__file__).parent / "test_data"
TEST_DATA_DIR.mkdir(exist_ok=True)

# Sample document paths
SAMPLE_DOCX = TEST_DATA_DIR / "sample.docx"
SAMPLE_PDF = TEST_DATA_DIR / "sample.pdf"

# Sample video URL (using a public test video)
SAMPLE_VIDEO_URL = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"

# Test credentials
TEST_USERNAME = "test_user"
TEST_PASSWORD = "test_pass"

# Test note structure
SAMPLE_NOTES = {
    "main_topics": [
        {
            "topic": "Test Topic",
            "subtopics": [
                {
                    "subtopic": "Test Subtopic",
                    "points": ["Test Point 1", "Test Point 2"]
                }
            ]
        }
    ]
}

# Create test documents if they don't exist
def create_test_documents():
    """Create test documents if they don't exist."""
    if not SAMPLE_DOCX.exists():
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.add_paragraph("For testing purposes.")
        doc.save(SAMPLE_DOCX)

    if not SAMPLE_PDF.exists():
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(SAMPLE_PDF))
        c.drawString(100, 750, "This is a test PDF document.")
        c.drawString(100, 700, "It contains multiple lines.")
 